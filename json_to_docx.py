import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def apply_font_style(run, style):
    if not style:
        return
    font = run.font
    font_name = style.get('font_name')
    if font_name:
        font.name = font_name
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), font_name)
    font_size_pt = style.get('font_size_pt')
    if font_size_pt:
        font.size = Pt(font_size_pt)

    font_weight = style.get('font_weight', '').lower()
    if font_weight == 'bold':
        font.bold = True
    else:
        font.bold = False

    if style.get('italic'):
        font.italic = True
    else:
        font.italic = False

    if style.get('underline'):
        font.underline = True
    else:
        font.underline = False

    color_hex = style.get('color_hex')
    if color_hex:
        color_hex = color_hex.lstrip('#')
        if len(color_hex) == 6:
            r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
            font.color.rgb = RGBColor(r, g, b)


def set_paragraph_alignment(paragraph, alignment):
    if not alignment:
        return
    mapping = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }
    align_value = mapping.get(alignment.lower())
    if align_value:
        paragraph.alignment = align_value


def set_paragraph_spacing(paragraph, style):
    # Новата функция добавена за line spacing и разстояния преди и след параграфа
    if not style:
        return
    p_format = paragraph.paragraph_format
    line_spacing_mult = style.get('line_spacing_mult')
    if line_spacing_mult:
        p_format.line_spacing = float(line_spacing_mult)
    space_before_pt = style.get('space_before_pt')
    if space_before_pt:
        p_format.space_before = Pt(space_before_pt)
    space_after_pt = style.get('space_after_pt')
    if space_after_pt:
        p_format.space_after = Pt(space_after_pt)


def set_cell_border_color(cell, border_color_hex, border_width_pt):
    # Добавя border със зададен цвят и дебелина към клетка (пример за таблици)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.findall(qn('w:tcBorders'))
    if not borders:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    else:
        borders = borders[0]

    def create_border_element(name):
        border = OxmlElement(f'w:{name}')
        if border_color_hex:
            border.set(qn('w:color'), border_color_hex.lstrip('#'))
        if border_width_pt:
            # Word използва 8ths от point, така че множим по 8
            border.set(qn('w:sz'), str(int(border_width_pt * 8)))
        border.set(qn('w:val'), 'single')
        return border

    for edge in ['top', 'left', 'bottom', 'right']:
        border_element = create_border_element(edge)
        borders.append(border_element)


def create_docx_from_json(data):
    document = Document()
    sections = data.get('document', {}).get('sections', [])
    for section in sections:
        section_type = section.get('type')
        if section_type == 'paragraph':
            para = document.add_paragraph()
            run = para.add_run(section.get('text', ''))
            style = section.get('style', {})
            apply_font_style(run, style)
            set_paragraph_alignment(para, style.get('alignment'))
            set_paragraph_spacing(para, style)
        elif section_type == 'table':
            rows_content = section.get('rows_content', [])
            # Автоматично определяне на брой редове и колони по съдържанието, с безопасна стойност
            rows = len(rows_content)
            columns = max((len(row) for row in rows_content), default=0)

            if rows > 0 and columns > 0:
                table = document.add_table(rows=rows, cols=columns)

                cell_style = section.get('cell_style', {})
                borders = section.get('borders', {})
                for r_idx, row_content in enumerate(rows_content):
                    for c_idx, cell_text in enumerate(row_content):
                        if r_idx < rows and c_idx < columns:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            paragraphs = cell.paragraphs
                            if paragraphs:
                                run = paragraphs[0].runs[0] if paragraphs[0].runs else paragraphs[0].add_run()
                                apply_font_style(run, cell_style)
                                set_paragraph_alignment(paragraphs[0], cell_style.get('alignment'))
                            if borders:
                                outer = borders.get('outer', {})
                                width_pt = outer.get('width_pt', 0)
                                color_hex = outer.get('color_hex')
                                set_cell_border_color(cell, color_hex, width_pt)
        elif section_type == 'image':
            description = section.get('description', '')
            para = document.add_paragraph(description)
        else:
            pass
    return document


def main():
    if len(sys.argv) < 2:
        print("Usage: python json_to_docx.py input.json")
        return

    json_file = sys.argv[1]
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = create_docx_from_json(data)
    output_file = json_file.rsplit('.', 1)[0] + '.docx'
    doc.save(output_file)
    print(f"Saved DOCX file as {output_file}")

if __name__ == "__main__":
    main()