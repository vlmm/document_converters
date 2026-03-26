"""
Python модул за конвертиране на Markdown в DOCX формат
Поддържа използване от командния ред
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import argparse
import os
from pathlib import Path
from typing import List, Tuple


_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$"
)

# Border width (in eighths of a point), spacing, and default color for table borders.
_TABLE_BORDER_SZ = 4
_TABLE_BORDER_SPACE = 0
_TABLE_BORDER_COLOR = "808080"  # RGB 128, 128, 128


def _parse_color(color_str: str) -> str:
    """Parse *color_str* to a 6-character uppercase hex string (e.g. ``'808080'``).

    Accepted formats:
    - Hex string: ``'808080'`` or ``'#808080'``
    - RGB triple: ``'128,128,128'``
    """
    s = color_str.strip().lstrip('#')
    if ',' in s:
        parts = s.split(',')
        if len(parts) != 3:
            raise ValueError(f"Invalid RGB color '{color_str}': expected 3 comma-separated integers")
        r, g, b = (int(p.strip()) for p in parts)
        for name, val in (('R', r), ('G', g), ('B', b)):
            if not 0 <= val <= 255:
                raise ValueError(f"RGB component {name}={val} out of range 0-255")
        return f"{r:02X}{g:02X}{b:02X}"
    if len(s) == 6 and all(c in '0123456789abcdefABCDEF' for c in s):
        return s.upper()
    raise ValueError(
        f"Invalid color '{color_str}': use a 6-digit hex string or 'R,G,B' integers"
    )


class MarkdownToDocx:
    def __init__(self, table_borders: bool = True, table_border_color: str = _TABLE_BORDER_COLOR):
        self.doc = Document()
        self.in_code_block = False
        self.code_block_lines = []
        self.table_buffer: List[str] = []
        self.table_borders = table_borders
        self.table_border_color = _parse_color(table_border_color)

    def convert(self, markdown_text: str) -> Document:
        """Конвертира markdown текст в DOCX документ"""
        lines = markdown_text.split('\n')

        for line in lines:
            self._process_line(line)

        # Flush pending table and code block at EOF
        self._flush_table()
        if self.in_code_block:
            self._add_code_block()

        return self.doc

    def _process_line(self, line: str):
        """Обработва един ред от markdown"""

        # Code блокове (highest priority – flush table first)
        if line.strip().startswith('```'):
            self._flush_table()
            if not self.in_code_block:
                self.in_code_block = True
                self.code_block_lines = []
            else:
                self._add_code_block()
                self.in_code_block = False
            return

        if self.in_code_block:
            self.code_block_lines.append(line)
            return

        # Таблични редове – буферират се
        if self._is_table_line(line):
            self.table_buffer.append(line)
            return

        # Не-таблично съдържание – изтриваме буфера
        self._flush_table()

        # Пропускане на празни редове
        if not line.strip():
            return

        # Заглавия
        if line.startswith('# '):
            self._add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            self._add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            self._add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            self._add_heading(line[5:].strip(), level=4)

        # Списъци
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            self._add_list_item(line.strip()[2:].strip())
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^\d+\.\s(.+)', line.strip())
            if match:
                self._add_list_item(match.group(1), ordered=True)

        # Обикновен текст с форматиране
        else:
            self._add_paragraph_with_formatting(line)

    # ------------------------------------------------------------------
    # Table helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _is_table_line(line: str) -> bool:
        """Return True if *line* looks like part of a Markdown pipe table."""
        s = line.strip()
        if '|' not in s:
            return False
        if bool(_TABLE_SEPARATOR_RE.match(s)):
            return True
        if s.startswith('|') or s.endswith('|'):
            return True
        # A single pipe is enough to detect tables without outer delimiters
        # (e.g. "A | B" produces a 2-column table with one pipe).
        # The fallback in _flush_table returns early and renders lines as
        # ordinary paragraphs when the second buffered line is not a valid
        # separator, so false positives (prose containing a lone "|") are safe.
        return s.count('|') >= 1

    @staticmethod
    def _is_separator_line(line: str) -> bool:
        return bool(_TABLE_SEPARATOR_RE.match(line.strip()))

    @staticmethod
    def _split_table_cells(line: str) -> List[str]:
        """Split a table row into individual cell strings.

        Handles:
        - leading/trailing pipes
        - escaped pipes (``\\|``) inside cell content
        """
        s = line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|') and not s.endswith('\\|'):
            s = s[:-1]
        # Split on unescaped pipes
        cells = re.split(r'(?<!\\)\|', s)
        return [c.strip().replace('\\|', '|') for c in cells]

    @staticmethod
    def _parse_alignments(separator_line: str) -> List[str]:
        """Return a list of alignment strings ('left', 'right', 'center')
        derived from a Markdown table separator row."""
        s = separator_line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|'):
            s = s[:-1]
        alignments = []
        for cell in s.split('|'):
            cell = cell.strip()
            if cell.startswith(':') and cell.endswith(':'):
                alignments.append('center')
            elif cell.endswith(':'):
                alignments.append('right')
            else:
                alignments.append('left')
        return alignments

    def _flush_table(self):
        """Convert any buffered table lines into a DOCX table, then clear the buffer."""
        if not self.table_buffer:
            return

        lines = self.table_buffer
        self.table_buffer = []

        # Validate: need header + separator as the first two lines
        if len(lines) < 2 or not self._is_separator_line(lines[1]):
            # Not a valid Markdown table – render as ordinary paragraphs
            for line in lines:
                if line.strip():
                    self._add_paragraph_with_formatting(line)
            return

        self._add_docx_table(lines)

    def _add_docx_table(self, lines: List[str]):
        """Build a python-docx table from a list of raw Markdown table lines."""
        alignments = self._parse_alignments(lines[1])

        # Data rows: skip the separator (index 1)
        data_lines = [lines[0]] + lines[2:]
        all_rows = [self._split_table_cells(l) for l in data_lines]

        if not all_rows:
            return

        num_cols = max(len(row) for row in all_rows)
        if num_cols == 0:
            return

        # Pad alignments list to num_cols
        while len(alignments) < num_cols:
            alignments.append('left')

        _align_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        }

        table = self.doc.add_table(rows=len(all_rows), cols=num_cols)
        self._remove_preferred_widths(table)
        self._set_table_borders(table, self.table_borders, self.table_border_color)

        for r_idx, row_cells in enumerate(all_rows):
            # Pad short rows
            while len(row_cells) < num_cols:
                row_cells.append('')
            for c_idx in range(num_cols):
                cell = table.cell(r_idx, c_idx)
                para = cell.paragraphs[0]
                run = para.add_run(row_cells[c_idx])
                if r_idx == 0:
                    run.bold = True
                para.alignment = _align_map.get(
                    alignments[c_idx] if c_idx < len(alignments) else 'left',
                    WD_PARAGRAPH_ALIGNMENT.LEFT,
                )

    @staticmethod
    def _set_table_borders(table, visible: bool, color: str = _TABLE_BORDER_COLOR):
        """Apply single-line borders in *color* (visible=True) or no borders (visible=False)."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn

        ns = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        if visible:
            b = (f'w:val="single" w:sz="{_TABLE_BORDER_SZ}" '
                 f'w:space="{_TABLE_BORDER_SPACE}" w:color="{color}"')
        else:
            b = (f'w:val="none" w:sz="0" '
                 f'w:space="{_TABLE_BORDER_SPACE}" w:color="auto"')

        borders_xml = (
            f'<w:tblBorders {ns}>'
            f'<w:top {b}/><w:left {b}/><w:bottom {b}/>'
            f'<w:right {b}/><w:insideH {b}/><w:insideV {b}/>'
            f'</w:tblBorders>'
        )
        borders_elm = parse_xml(borders_xml)

        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {ns}/>')
            table._element.insert(0, tblPr)

        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(borders_elm)

    @staticmethod
    def _remove_preferred_widths(table):
        """Remove table-level, column-level and cell-level preferred-width settings.

        python-docx may write ``w:tblW`` on the table properties, ``w:w`` on
        each ``w:gridCol``, and ``w:tcW`` on each cell's ``w:tcPr``.  Removing
        these lets Word calculate column widths automatically.
        """
        from docx.oxml.ns import qn

        # Table preferred width
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblPr.remove(tblW)

        # Column preferred widths in tblGrid
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gridCol in tblGrid.findall(qn('w:gridCol')):
                attrib_key = qn('w:w')
                if attrib_key in gridCol.attrib:
                    del gridCol.attrib[attrib_key]

        # Cell preferred widths
        for tc in table._element.iter(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)

    def _add_heading(self, text: str, level: int):
        """Добавя заглавие"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _add_list_item(self, text: str, ordered: bool = False):
        """Добавя елемент на списък"""
        paragraph = self.doc.add_paragraph(text, style='List Number' if ordered else 'List Bullet')

    def _add_code_block(self):
        """Добавя code блок"""
        code_text = '\n'.join(self.code_block_lines)
        paragraph = self.doc.add_paragraph(code_text)

        # Форматиране на код
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Светлосиво фоново оцветяване
        shading_elm = self._shade_paragraph(paragraph)

    def _add_paragraph_with_formatting(self, text: str):
        """Добавя параграф с bold, italic, strikethrough и inline код форматиране"""
        paragraph = self.doc.add_paragraph()

        # Обработка на форматирането в правилния ред (от най-специфично към по-общо)
        tokens = self._tokenize_formatting(text)

        for token_type, token_text in tokens:
            if not token_text:  # Пропускане на празни токени
                continue

            run = paragraph.add_run(token_text)

            if token_type == 'bold':
                run.bold = True
            elif token_type == 'italic':
                run.italic = True
            elif token_type == 'strikethrough':
                run.font.strike = True
            elif token_type == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(192, 0, 0)

    def _tokenize_formatting(self, text: str):
        """Разделя текста на токени с информация за форматирането"""
        tokens = []
        i = 0

        while i < len(text):
            # Inline код - приоритет 1 (най-висок)
            if i < len(text) - 1 and text[i] == '`':
                end = text.find('`', i + 1)
                if end != -1:
                    tokens.append(('code', text[i+1:end]))
                    i = end + 1
                    continue

            # Bold (**текст**) - приоритет 2
            if i + 1 < len(text) and text[i:i+2] == '**':
                end = text.find('**', i + 2)
                if end != -1:
                    tokens.append(('bold', text[i+2:end]))
                    i = end + 2
                    continue

            # Зачеркнат текст (~~текст~~) - приоритет 3
            if i + 1 < len(text) and text[i:i+2] == '~~':
                end = text.find('~~', i + 2)
                if end != -1:
                    tokens.append(('strikethrough', text[i+2:end]))
                    i = end + 2
                    continue

            # Italic (*текст* или _текст_) - приоритет 4
            if text[i] in ('*', '_'):
                marker = text[i]
                end = text.find(marker, i + 1)
                if end != -1:
                    # Проверка да не е част от **
                    if marker == '*' and i > 0 and text[i-1] == '*':
                        # Вече е обработено в bold
                        tokens.append(('normal', text[i]))
                        i += 1
                        continue
                    if marker == '*' and end + 1 < len(text) and text[end+1] == '*':
                        # Следва второ * - част от **
                        tokens.append(('normal', text[i]))
                        i += 1
                        continue

                    tokens.append(('italic', text[i+1:end]))
                    i = end + 1
                    continue

            # Обикновен текст - намиране на следващия специален символ
            next_special = len(text)
            patterns = ['`', '**', '~~', '*', '_']

            for pattern in patterns:
                pos = text.find(pattern, i)
                if pos != -1 and pos < next_special:
                    next_special = pos

            if next_special > i:
                tokens.append(('normal', text[i:next_special]))
                i = next_special
            else:
                # Остатък от текста
                tokens.append(('normal', text[i:]))
                break

        return tokens

    def _shade_paragraph(self, paragraph):
        """Добавя сива фонова оцветяване на параграф"""
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        paragraph._element.get_or_add_pPr().append(shading_elm)
        return shading_elm

    def save(self, filename: str):
        """Запазва документа като DOCX файл"""
        self.doc.save(filename)


def markdown_to_docx(markdown_text: str, output_filename: str,
                     table_borders: bool = True,
                     table_border_color: str = _TABLE_BORDER_COLOR):
    """Простна функция за конвертиране на markdown в docx"""
    converter = MarkdownToDocx(table_borders=table_borders, table_border_color=table_border_color)
    doc = converter.convert(markdown_text)
    converter.save(output_filename)


def convert_file(input_file: str, output_file: str = None,
                 table_borders: bool = True,
                 table_border_color: str = _TABLE_BORDER_COLOR):
    """Конвертира markdown файл в docx файл"""

    # Проверка дали входния файл съществува
    if not os.path.exists(input_file):
        print(f"❌ Грешка: Файлът '{input_file}' не съществува!")
        return False

    # Проверка дали файлът е markdown
    if not input_file.lower().endswith(('.md', '.markdown')):
        print(f"⚠️  Внимание: Файлът не е markdown формат")

    # Генериране на изходния файл ако не е посочен
    if output_file is None:
        base_name = Path(input_file).stem
        output_file = f"{base_name}.docx"

    try:
        # Прочитане на markdown файла
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        print(f"📖 Прочитане на файл: {input_file}")

        # Конвертиране
        print(f"⏳ Конвертиране на markdown в docx...")
        markdown_to_docx(markdown_content, output_file,
                         table_borders=table_borders,
                         table_border_color=table_border_color)

        print(f"✅ Успешно! Документ създаден: {output_file}")
        return True

    except Exception as e:
        print(f"❌ Грешка при конвертиране: {str(e)}")
        return False


def main():
    """Главна функция за командния ред"""
    parser = argparse.ArgumentParser(
        description='Конвертира Markdown файлове в DOCX формат',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примери за използване:
  python markdown_to_docx.py README.md
  python markdown_to_docx.py README.md -o output.docx
  python markdown_to_docx.py document.md --output report.docx
  python markdown_to_docx.py document.md --no-table-borders
  python markdown_to_docx.py document.md --table-border-color 000000
  python markdown_to_docx.py document.md --table-border-color "128,128,128"

Поддържани форматирания:
  **bold текст** - удебелен текст
  ~~strikethrough~~ - зачеркнат текст
  *italic текст* - наклонен текст
  `inline код` - вътрешен код
        """
    )

    parser.add_argument(
        'input',
        help='Пътя до markdown файла за конвертиране'
    )

    parser.add_argument(
        '-o', '--output',
        help='Пътя до изходния DOCX файл (по подразбиране се използва същото име с .docx разширение)',
        metavar='OUTPUT_FILE'
    )

    parser.add_argument(
        '-v', '--version',
        action='version',
        version='%(prog)s 1.0'
    )

    parser.add_argument(
        '--no-table-borders',
        action='store_true',
        default=False,
        help='Скрива рамките на таблиците (по подразбиране рамките са видими)',
    )

    parser.add_argument(
        '--table-border-color',
        default=_TABLE_BORDER_COLOR,
        metavar='COLOR',
        help=(
            'Цвят на рамките на таблиците – 6-цифрен hex (напр. 808080) '
            'или R,G,B цели числа (напр. "128,128,128"). '
            f'По подразбиране: {_TABLE_BORDER_COLOR} (RGB 128,128,128)'
        ),
    )

    args = parser.parse_args()

    try:
        border_color = _parse_color(args.table_border_color)
    except ValueError as exc:
        parser.error(str(exc))

    # Конвертиране на файла
    success = convert_file(
        args.input,
        args.output,
        table_borders=not args.no_table_borders,
        table_border_color=border_color,
    )

    # Връщане на статус код
    return 0 if success else 1


if __name__ == "__main__":
    exit(main())