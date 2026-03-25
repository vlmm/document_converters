"""
Tests for docx_to_markdown.py

Covers:
1. Nested/overlapping inline styles (bold+italic, bold+underline, bold+italic+strike, etc.)
2. Bullet and numbered lists with 2+ nesting levels (style-name detection + XML numPr)
3. Indented non-list paragraphs
4. Tables with:
   - multi-paragraph cells (joined with <br>)
   - list paragraphs inside cells
   - indentation converted to &nbsp;
"""

import io
import sys
import os
import unittest
from pathlib import Path
from unittest.mock import MagicMock

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx_to_markdown import (
    _runs_to_markdown_text,
    _get_list_info,
    _get_paragraph_indent_level,
    _paragraph_to_md_line,
    _convert_cell_to_md,
    _table_row_cells,
    docx_to_markdown,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_run(text: str, bold=False, italic=False, underline=False, strike=False, color_rgb=None):
    """Create a minimal mock run object.

    ``color_rgb`` may be a hex string like ``"FF0000"`` for an explicit color,
    or ``None`` (default) to represent the automatic/default black color.
    """
    run = MagicMock()
    run.text = text
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font = MagicMock()
    run.font.strike = strike
    run.font.color = MagicMock()
    run.font.color.rgb = color_rgb
    # .type mirrors the python-docx ColorFormat API (None = auto/inherit).
    run.font.color.type = None if color_rgb is None else "explicit"
    return run


def _make_paragraph(
    runs,
    style_name: str = "Normal",
    left_indent=None,
    first_line_indent=None,
    num_id: int = 0,
    ilvl: int = 0,
):
    """Create a minimal mock paragraph."""
    p = MagicMock()
    p.runs = runs

    style = MagicMock()
    style.name = style_name
    p.style = style

    pf = MagicMock()
    pf.left_indent = left_indent
    pf.first_line_indent = first_line_indent
    p.paragraph_format = pf

    # Build XML-like numPr structure
    if num_id > 0:
        numId_elem = MagicMock()
        numId_elem.val = num_id
        ilvl_elem = MagicMock()
        ilvl_elem.val = ilvl
        numPr = MagicMock()
        numPr.numId = numId_elem
        numPr.ilvl = ilvl_elem
        pPr = MagicMock()
        pPr.numPr = numPr
        p._p = MagicMock()
        p._p.pPr = pPr
        p._p.xpath.return_value = []
    else:
        # no numPr
        pPr = MagicMock()
        pPr.numPr = None
        p._p = MagicMock()
        p._p.pPr = pPr
        p._p.xpath.return_value = []

    return p


# ---------------------------------------------------------------------------
# 1. Inline formatting
# ---------------------------------------------------------------------------

class TestRunsToMarkdownText(unittest.TestCase):

    def test_plain_text(self):
        runs = [_make_run("hello")]
        self.assertEqual(_runs_to_markdown_text(runs), "hello")

    def test_bold_only(self):
        runs = [_make_run("bold", bold=True)]
        self.assertEqual(_runs_to_markdown_text(runs), "**bold**")

    def test_italic_only(self):
        runs = [_make_run("ital", italic=True)]
        self.assertEqual(_runs_to_markdown_text(runs), "_ital_")

    def test_underline_only(self):
        runs = [_make_run("ul", underline=True)]
        self.assertEqual(_runs_to_markdown_text(runs), "<u>ul</u>")

    def test_strike_only(self):
        runs = [_make_run("st", strike=True)]
        self.assertEqual(_runs_to_markdown_text(runs), "~~st~~")

    def test_bold_italic_same_run(self):
        runs = [_make_run("bi", bold=True, italic=True)]
        result = _runs_to_markdown_text(runs)
        # Both markers must be present and properly closed.
        self.assertIn("**", result)
        self.assertIn("_", result)
        self.assertIn("bi", result)
        # Properly nested: last-opened closes first.
        self.assertTrue(
            result.endswith("_**") or result.endswith("**_"),
            f"Expected nested closing, got: {result!r}",
        )

    def test_bold_then_italic_separate_runs(self):
        """bold text followed by italic text – no mis-nesting."""
        runs = [_make_run("A", bold=True), _make_run("B", italic=True)]
        result = _runs_to_markdown_text(runs)
        # Bold must be closed before italic opens.
        self.assertEqual(result, "**A**_B_")

    def test_bold_italic_to_bold_only(self):
        """Transition from bold+italic to bold-only must close italic only."""
        runs = [
            _make_run("bi", bold=True, italic=True),
            _make_run("b", bold=True),
        ]
        result = _runs_to_markdown_text(runs)
        # Expected: **_bi_b**  (italic closes, bold stays, then closes at end)
        self.assertEqual(result, "**_bi_b**")

    def test_bold_italic_underline_all_three(self):
        runs = [_make_run("all", bold=True, italic=True, underline=True)]
        result = _runs_to_markdown_text(runs)
        self.assertIn("**", result)
        self.assertIn("_", result)
        self.assertIn("<u>", result)
        self.assertIn("</u>", result)
        self.assertIn("all", result)

    def test_style_transition_italic_to_bold_italic(self):
        """italic → bold+italic: bold must be opened without breaking italic."""
        runs = [
            _make_run("i", italic=True),
            _make_run("bi", bold=True, italic=True),
        ]
        result = _runs_to_markdown_text(runs)
        # Expected: _i_ then **_bi_**  (italic closes, bold+italic opens)
        # or _i**bi**_ if italic stays open and bold opens inside
        # Because italic is common prefix for the second run but order changes,
        # we check structural correctness:
        self.assertIn("i", result)
        self.assertIn("bi", result)
        # Must be properly closed (no dangling markers)
        self.assertEqual(result.count("**"), 2)
        self.assertEqual(result.count("_") % 2, 0)  # even number of _

    def test_empty_runs_ignored(self):
        runs = [_make_run(""), _make_run("x", bold=True), _make_run("")]
        self.assertEqual(_runs_to_markdown_text(runs), "**x**")

    def test_no_runs(self):
        self.assertEqual(_runs_to_markdown_text([]), "")

    def test_bold_strike_combination(self):
        runs = [_make_run("bs", bold=True, strike=True)]
        result = _runs_to_markdown_text(runs)
        self.assertIn("**", result)
        self.assertIn("~~", result)
        self.assertIn("bs", result)

    def test_multiple_runs_no_style_change(self):
        """Consecutive runs with same style should not add extra markers."""
        runs = [_make_run("A", bold=True), _make_run("B", bold=True)]
        self.assertEqual(_runs_to_markdown_text(runs), "**AB**")

    def test_mixed_plain_and_bold(self):
        runs = [_make_run("plain "), _make_run("bold", bold=True), _make_run(" plain")]
        self.assertEqual(_runs_to_markdown_text(runs), "plain **bold** plain")


# ---------------------------------------------------------------------------
# 2. List detection and nesting
# ---------------------------------------------------------------------------

class TestGetListInfo(unittest.TestCase):

    def test_bullet_style_no_xml(self):
        p = _make_paragraph([], style_name="List Bullet")
        is_list, level, is_numbered = _get_list_info(p)
        self.assertTrue(is_list)
        self.assertFalse(is_numbered)

    def test_number_style_no_xml(self):
        p = _make_paragraph([], style_name="List Number")
        is_list, level, is_numbered = _get_list_info(p)
        self.assertTrue(is_list)
        self.assertTrue(is_numbered)

    def test_normal_style_not_list(self):
        p = _make_paragraph([], style_name="Normal")
        is_list, level, is_numbered = _get_list_info(p)
        self.assertFalse(is_list)

    def test_xml_numpr_level_0(self):
        p = _make_paragraph([], style_name="List Bullet", num_id=1, ilvl=0)
        is_list, level, is_numbered = _get_list_info(p)
        self.assertTrue(is_list)
        self.assertEqual(level, 0)

    def test_xml_numpr_level_1(self):
        p = _make_paragraph([], style_name="List Bullet", num_id=1, ilvl=1)
        is_list, level, is_numbered = _get_list_info(p)
        self.assertTrue(is_list)
        self.assertEqual(level, 1)

    def test_xml_numpr_level_2(self):
        p = _make_paragraph([], style_name="List Bullet", num_id=1, ilvl=2)
        is_list, level, is_numbered = _get_list_info(p)
        self.assertTrue(is_list)
        self.assertEqual(level, 2)

    def test_xml_numid_zero_not_list(self):
        """numId=0 means no active numbering – not a list."""
        p = _make_paragraph([], style_name="Normal", num_id=0)
        is_list, level, is_numbered = _get_list_info(p)
        self.assertFalse(is_list)


class TestParagraphToMdLine(unittest.TestCase):

    def test_bullet_level0(self):
        runs = [_make_run("Item")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=0)
        self.assertEqual(_paragraph_to_md_line(p), "- Item")

    def test_bullet_level1(self):
        runs = [_make_run("Sub")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=1)
        self.assertEqual(_paragraph_to_md_line(p), "  - Sub")

    def test_bullet_level2(self):
        runs = [_make_run("SubSub")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=2)
        self.assertEqual(_paragraph_to_md_line(p), "    - SubSub")

    def test_numbered_level0(self):
        runs = [_make_run("First")]
        p = _make_paragraph(runs, style_name="List Number", num_id=2, ilvl=0)
        self.assertEqual(_paragraph_to_md_line(p), "1. First")

    def test_numbered_level1(self):
        runs = [_make_run("Second")]
        p = _make_paragraph(runs, style_name="List Number", num_id=2, ilvl=1)
        self.assertEqual(_paragraph_to_md_line(p), "  1. Second")

    def test_letter_clause_not_list(self):
        """(a) clause paragraphs must not become list items."""
        runs = [_make_run("(a) some clause")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=0)
        result = _paragraph_to_md_line(p)
        self.assertFalse(result.startswith("-"), f"Should not be a list: {result!r}")
        self.assertIn("(a) some clause", result)

    def test_empty_paragraph(self):
        p = _make_paragraph([_make_run("")])
        self.assertEqual(_paragraph_to_md_line(p), "")


# ---------------------------------------------------------------------------
# 3. Indented non-list paragraphs
# ---------------------------------------------------------------------------

class TestIndentation(unittest.TestCase):

    # 457200 EMU = 0.5 inch = level 1; 914400 = level 2
    _STEP = 457_200

    def test_no_indent(self):
        runs = [_make_run("plain")]
        p = _make_paragraph(runs, left_indent=None)
        self.assertEqual(_get_paragraph_indent_level(p), 0)
        self.assertEqual(_paragraph_to_md_line(p), "plain")

    def test_indent_level1(self):
        runs = [_make_run("indented")]
        p = _make_paragraph(runs, left_indent=self._STEP)
        self.assertEqual(_get_paragraph_indent_level(p), 1)
        self.assertEqual(_paragraph_to_md_line(p), "  indented")

    def test_indent_level2(self):
        runs = [_make_run("deep")]
        p = _make_paragraph(runs, left_indent=2 * self._STEP)
        self.assertEqual(_get_paragraph_indent_level(p), 2)
        self.assertEqual(_paragraph_to_md_line(p), "    deep")

    def test_indent_level3(self):
        runs = [_make_run("deeper")]
        p = _make_paragraph(runs, left_indent=3 * self._STEP)
        self.assertEqual(_get_paragraph_indent_level(p), 3)
        self.assertEqual(_paragraph_to_md_line(p), "      deeper")

    def test_first_line_indent_used_when_larger(self):
        runs = [_make_run("fli")]
        p = _make_paragraph(runs, left_indent=None, first_line_indent=self._STEP)
        self.assertEqual(_get_paragraph_indent_level(p), 1)

    def test_indent_in_table_uses_nbsp(self):
        runs = [_make_run("cell")]
        p = _make_paragraph(runs, left_indent=self._STEP)
        result = _paragraph_to_md_line(p, in_table=True)
        self.assertTrue(result.startswith("&nbsp;"), f"Expected &nbsp;, got: {result!r}")

    def test_list_ignores_indent_level(self):
        """List paragraphs use ilvl for nesting, not indent_level."""
        runs = [_make_run("item")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=0,
                            left_indent=self._STEP)
        result = _paragraph_to_md_line(p)
        # Should be a list item with ilvl=0 (no extra indent)
        self.assertEqual(result, "- item")


# ---------------------------------------------------------------------------
# 4. Table cell conversion
# ---------------------------------------------------------------------------

class TestConvertCellToMd(unittest.TestCase):

    def _make_cell(self, paragraphs):
        cell = MagicMock()
        cell.paragraphs = paragraphs
        return cell

    def test_single_paragraph(self):
        runs = [_make_run("hello")]
        p = _make_paragraph(runs)
        cell = self._make_cell([p])
        self.assertEqual(_convert_cell_to_md(cell), "hello")

    def test_multi_paragraph_joined_with_br(self):
        p1 = _make_paragraph([_make_run("line1")])
        p2 = _make_paragraph([_make_run("line2")])
        cell = self._make_cell([p1, p2])
        self.assertEqual(_convert_cell_to_md(cell), "line1<br>line2")

    def test_three_paragraphs(self):
        paras = [_make_paragraph([_make_run(f"p{i}")]) for i in range(3)]
        cell = self._make_cell(paras)
        self.assertEqual(_convert_cell_to_md(cell), "p0<br>p1<br>p2")

    def test_empty_paragraphs_skipped(self):
        p1 = _make_paragraph([_make_run("A")])
        p_empty = _make_paragraph([_make_run("")])
        p2 = _make_paragraph([_make_run("B")])
        cell = self._make_cell([p1, p_empty, p2])
        self.assertEqual(_convert_cell_to_md(cell), "A<br>B")

    def test_list_in_cell(self):
        runs = [_make_run("item")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=0)
        cell = self._make_cell([p])
        result = _convert_cell_to_md(cell)
        self.assertIn("-", result)
        self.assertIn("item", result)

    def test_nested_list_in_cell(self):
        runs = [_make_run("sub")]
        p = _make_paragraph(runs, style_name="List Bullet", num_id=1, ilvl=1)
        cell = self._make_cell([p])
        result = _convert_cell_to_md(cell)
        # Leading spaces from indentation should be turned into &nbsp;
        self.assertIn("&nbsp;", result)
        self.assertIn("-", result)
        self.assertIn("sub", result)

    def test_indent_in_cell_uses_nbsp(self):
        _STEP = 457_200
        runs = [_make_run("indented")]
        p = _make_paragraph(runs, left_indent=_STEP)
        cell = self._make_cell([p])
        result = _convert_cell_to_md(cell)
        self.assertIn("&nbsp;", result)

    def test_empty_cell(self):
        p = _make_paragraph([_make_run("")])
        cell = self._make_cell([p])
        self.assertEqual(_convert_cell_to_md(cell), "")


# ---------------------------------------------------------------------------
# 5. End-to-end: convert real .docx documents
# ---------------------------------------------------------------------------

try:
    from docx import Document as DocxDocument  # type: ignore
    from docx.shared import Pt, Inches  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from lxml import etree  # type: ignore
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _add_list_paragraph(doc, text, style_name="List Bullet", level=0):
    """Add a list paragraph with proper numbering XML."""
    p = doc.add_paragraph(text, style=style_name)
    # Set ilvl via XML
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = level
    return p


@unittest.skipUnless(_DOCX_AVAILABLE, "python-docx not installed")
class TestEndToEnd(unittest.TestCase):

    def _convert(self, doc):
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = Path(f.name)
        try:
            doc.save(str(tmp))
            return docx_to_markdown(tmp)
        finally:
            tmp.unlink(missing_ok=True)

    # -- Inline formatting --

    def test_e2e_bold(self):
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("bold").bold = True
        result = self._convert(doc)
        self.assertIn("**bold**", result)

    def test_e2e_italic(self):
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("ital").italic = True
        result = self._convert(doc)
        self.assertIn("_ital_", result)

    def test_e2e_bold_italic_combined(self):
        doc = DocxDocument()
        p = doc.add_paragraph()
        r = p.add_run("bi")
        r.bold = True
        r.italic = True
        result = self._convert(doc)
        self.assertIn("**", result)
        self.assertIn("_", result)
        self.assertIn("bi", result)

    def test_e2e_bold_then_italic(self):
        """Bold run followed by italic run – properly closed and opened."""
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("A").bold = True
        p.add_run("B").italic = True
        result = self._convert(doc)
        self.assertIn("**A**", result)
        self.assertIn("_B_", result)

    # -- Lists --

    def test_e2e_bullet_list_level0(self):
        doc = DocxDocument()
        _add_list_paragraph(doc, "Item 1", "List Bullet", 0)
        _add_list_paragraph(doc, "Item 2", "List Bullet", 0)
        result = self._convert(doc)
        self.assertIn("- Item 1", result)
        self.assertIn("- Item 2", result)

    def test_e2e_bullet_list_level1(self):
        doc = DocxDocument()
        _add_list_paragraph(doc, "Parent", "List Bullet", 0)
        _add_list_paragraph(doc, "Child", "List Bullet", 1)
        result = self._convert(doc)
        self.assertIn("- Parent", result)
        self.assertIn("  - Child", result)

    def test_e2e_numbered_list(self):
        doc = DocxDocument()
        _add_list_paragraph(doc, "One", "List Number", 0)
        _add_list_paragraph(doc, "Two", "List Number", 0)
        result = self._convert(doc)
        self.assertIn("1. One", result)
        self.assertIn("1. Two", result)

    def test_e2e_letter_clause_not_list(self):
        doc = DocxDocument()
        p = doc.add_paragraph("(a) some clause", style="List Bullet")
        result = self._convert(doc)
        self.assertNotIn("- (a)", result)
        self.assertIn("(a) some clause", result)

    # -- Indentation --

    def test_e2e_indented_paragraph(self):
        doc = DocxDocument()
        p = doc.add_paragraph("indented text")
        p.paragraph_format.left_indent = Inches(0.5)  # level 1
        result = self._convert(doc)
        self.assertIn("  indented text", result)

    # -- Table --

    def test_e2e_table_basic(self):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].paragraphs[0].add_run("H1")
        table.rows[0].cells[1].paragraphs[0].add_run("H2")
        table.rows[1].cells[0].paragraphs[0].add_run("A")
        table.rows[1].cells[1].paragraphs[0].add_run("B")
        result = self._convert(doc)
        self.assertIn("| H1 | H2 |", result)
        self.assertIn("| A | B |", result)

    def test_e2e_table_multi_paragraph_cell(self):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=1)
        # Header
        table.rows[0].cells[0].paragraphs[0].add_run("Header")
        # Data cell with two paragraphs
        cell = table.rows[1].cells[0]
        cell.paragraphs[0].add_run("line1")
        cell.add_paragraph("line2")
        result = self._convert(doc)
        self.assertIn("<br>", result)
        self.assertIn("line1", result)
        self.assertIn("line2", result)

    def test_e2e_table_list_in_cell(self):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run("Header")
        cell = table.rows[1].cells[0]
        # Replace the default empty paragraph with a list paragraph
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run("bullet item")
        cell.paragraphs[0].style = doc.styles["List Bullet"]
        pPr = cell.paragraphs[0]._p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numPr.get_or_add_ilvl().val = 0
        numId = numPr.get_or_add_numId()
        numId.val = 1
        result = self._convert(doc)
        self.assertIn("- bullet item", result)

    def test_e2e_table_indentation_to_nbsp(self):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run("Header")
        cell = table.rows[1].cells[0]
        p = cell.paragraphs[0]
        p.add_run("indented")
        p.paragraph_format.left_indent = Inches(0.5)
        result = self._convert(doc)
        self.assertIn("&nbsp;", result)

    def test_e2e_table_horizontal_merge_no_duplicate_columns(self):
        """Horizontally merged cells must not produce duplicate columns."""
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=3)
        # Merge first two columns in header row
        table.rows[0].cells[0].merge(table.rows[0].cells[1])
        table.rows[0].cells[0].paragraphs[0].add_run("Merged")
        table.rows[0].cells[2].paragraphs[0].add_run("Col3")
        table.rows[1].cells[0].paragraphs[0].add_run("A")
        table.rows[1].cells[1].paragraphs[0].add_run("B")
        table.rows[1].cells[2].paragraphs[0].add_run("C")
        result = self._convert(doc)
        lines = result.splitlines()
        # Find the header row (contains "Merged")
        header_lines = [l for l in lines if "Merged" in l]
        self.assertEqual(len(header_lines), 1, f"Expected one header line, got: {header_lines}")
        header_line = header_lines[0]
        # After deduplication the merged cell appears once: 2 unique columns → 3 pipe chars
        col_count = header_line.count("|") - 1
        self.assertEqual(col_count, 2, f"Expected 2 columns after merge, got {col_count}: {header_line!r}")

    def test_e2e_italic_non_black(self):
        """--italic-non-black: colored runs are wrapped in _..._."""
        doc = DocxDocument()
        from docx.shared import RGBColor  # type: ignore
        p = doc.add_paragraph()
        r = p.add_run("colored")
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red
        result = self._convert_with_flag(doc, italic_non_black=True)
        self.assertIn("_colored_", result)

    def test_e2e_italic_non_black_black_unchanged(self):
        """--italic-non-black: explicit black text must NOT become italic."""
        doc = DocxDocument()
        from docx.shared import RGBColor  # type: ignore
        p = doc.add_paragraph()
        r = p.add_run("black")
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        result = self._convert_with_flag(doc, italic_non_black=True)
        self.assertNotIn("_black_", result)
        self.assertIn("black", result)

    def test_e2e_italic_non_black_auto_unchanged(self):
        """--italic-non-black: auto-color text (no explicit color) must NOT become italic."""
        doc = DocxDocument()
        p = doc.add_paragraph()
        p.add_run("auto")  # no color set → auto/black
        result = self._convert_with_flag(doc, italic_non_black=True)
        self.assertNotIn("_auto_", result)

    def test_e2e_italic_non_black_flag_off(self):
        """Without --italic-non-black, colored text is NOT italicized."""
        doc = DocxDocument()
        from docx.shared import RGBColor  # type: ignore
        p = doc.add_paragraph()
        r = p.add_run("colored")
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        result = self._convert_with_flag(doc, italic_non_black=False)
        self.assertNotIn("_colored_", result)
        self.assertIn("colored", result)

    def _convert_with_flag(self, doc, italic_non_black: bool = False):
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = Path(f.name)
        try:
            doc.save(str(tmp))
            return docx_to_markdown(tmp, italic_non_black=italic_non_black)
        finally:
            tmp.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# 6. _table_row_cells deduplication
# ---------------------------------------------------------------------------

class TestTableRowCells(unittest.TestCase):

    def _make_cell(self, tc):
        cell = MagicMock()
        cell._tc = tc
        return cell

    def test_deduplicates_horizontally_merged_cells(self):
        """row.cells returning the same cell twice must yield only one entry."""
        tc1 = MagicMock(name="tc1")
        tc2 = MagicMock(name="tc2")
        cell1 = self._make_cell(tc1)
        cell2 = self._make_cell(tc2)

        row = MagicMock()
        # cell1 repeated (horizontal merge spanning 2 columns)
        row.cells = (cell1, cell1, cell2)

        result = _table_row_cells(row)
        self.assertEqual(len(result), 2)
        self.assertIs(result[0], cell1)
        self.assertIs(result[1], cell2)

    def test_no_merge_returns_all_cells(self):
        """Rows without merges must return all cells unchanged."""
        tc1, tc2, tc3 = MagicMock(), MagicMock(), MagicMock()
        cell1 = self._make_cell(tc1)
        cell2 = self._make_cell(tc2)
        cell3 = self._make_cell(tc3)

        row = MagicMock()
        row.cells = (cell1, cell2, cell3)

        result = _table_row_cells(row)
        self.assertEqual(result, [cell1, cell2, cell3])

    def test_all_same_cell_returns_one(self):
        """Extreme case: entire row is one merged cell."""
        tc = MagicMock(name="tc")
        cell = self._make_cell(tc)

        row = MagicMock()
        row.cells = (cell, cell, cell)

        result = _table_row_cells(row)
        self.assertEqual(len(result), 1)
        self.assertIs(result[0], cell)


# ---------------------------------------------------------------------------
# 7. _is_non_black_color and italic_non_black flag
# ---------------------------------------------------------------------------

class TestItalicNonBlack(unittest.TestCase):

    def test_non_black_color_makes_italic(self):
        run = _make_run("red", color_rgb="FF0000")
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertIn("_red_", result)

    def test_explicit_black_not_italic(self):
        run = _make_run("black", color_rgb="000000")
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertEqual(result, "black")

    def test_auto_color_not_italic(self):
        run = _make_run("auto", color_rgb=None)
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertEqual(result, "auto")

    def test_flag_off_no_italic(self):
        run = _make_run("red", color_rgb="FF0000")
        result = _runs_to_markdown_text([run], italic_non_black=False)
        self.assertEqual(result, "red")

    def test_already_italic_plus_non_black_stays_italic(self):
        """Run that is already italic + non-black color: no double _."""
        run = _make_run("red+italic", italic=True, color_rgb="FF0000")
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertEqual(result, "_red+italic_")

    def test_non_black_combined_with_bold(self):
        """Non-black color + bold → bold italic."""
        run = _make_run("bi", bold=True, color_rgb="FF0000")
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertIn("**", result)
        self.assertIn("_", result)
        self.assertIn("bi", result)

    def test_color_exception_not_italic(self):
        """If color access raises, treat as black (no italic)."""
        from unittest.mock import PropertyMock
        run = _make_run("safe")
        run.font.color = MagicMock()
        # Make .rgb property itself raise when accessed
        type(run.font.color).rgb = PropertyMock(side_effect=Exception("no color"))
        result = _runs_to_markdown_text([run], italic_non_black=True)
        self.assertEqual(result, "safe")


if __name__ == "__main__":
    unittest.main()
