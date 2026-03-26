import sys
import os
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from markdown_to_docx import MarkdownToDocx, _parse_color

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def _convert(md: str):
    converter = MarkdownToDocx()
    return converter.convert(md)


def _tables(doc):
    return doc.tables


class TestMarkdownTableDetection(unittest.TestCase):
    """Tests for _is_table_line and _is_separator_line helpers."""

    def test_pipe_row_is_table_line(self):
        self.assertTrue(MarkdownToDocx._is_table_line("| A | B | C |"))

    def test_separator_is_table_line(self):
        self.assertTrue(MarkdownToDocx._is_table_line("| --- | --- | --- |"))

    def test_no_pipe_is_not_table_line(self):
        self.assertFalse(MarkdownToDocx._is_table_line("Normal text"))

    def test_row_without_outer_pipes_is_table_line(self):
        # Two or more pipes → table row
        self.assertTrue(MarkdownToDocx._is_table_line("A | B | C"))

    def test_separator_detected(self):
        self.assertTrue(MarkdownToDocx._is_separator_line("|---|---|"))
        self.assertTrue(MarkdownToDocx._is_separator_line("|:---|---:|:---:|"))
        self.assertFalse(MarkdownToDocx._is_separator_line("| A | B |"))


class TestSplitTableCells(unittest.TestCase):
    def test_basic_split(self):
        cells = MarkdownToDocx._split_table_cells("| A | B | C |")
        self.assertEqual(cells, ["A", "B", "C"])

    def test_no_outer_pipes(self):
        cells = MarkdownToDocx._split_table_cells("A | B | C")
        self.assertEqual(cells, ["A", "B", "C"])

    def test_escaped_pipe(self):
        cells = MarkdownToDocx._split_table_cells(r"| A\|B | C |")
        self.assertEqual(cells, ["A|B", "C"])

    def test_whitespace_trimmed(self):
        cells = MarkdownToDocx._split_table_cells("|  hello  |  world  |")
        self.assertEqual(cells, ["hello", "world"])


class TestParseAlignments(unittest.TestCase):
    def test_default_left(self):
        aligns = MarkdownToDocx._parse_alignments("|---|---|")
        self.assertEqual(aligns, ["left", "left"])

    def test_right_alignment(self):
        aligns = MarkdownToDocx._parse_alignments("|---|---:|")
        self.assertEqual(aligns, ["left", "right"])

    def test_center_alignment(self):
        aligns = MarkdownToDocx._parse_alignments("|:---:|:---:|")
        self.assertEqual(aligns, ["center", "center"])

    def test_mixed_alignments(self):
        aligns = MarkdownToDocx._parse_alignments("|:---|---:|:---:|")
        self.assertEqual(aligns, ["left", "right", "center"])


class TestBasicTableConversion(unittest.TestCase):
    SIMPLE_TABLE = "| Name | Age |\n|------|-----|\n| Alice | 30 |\n| Bob | 25 |"

    def test_table_is_created(self):
        doc = _convert(self.SIMPLE_TABLE)
        self.assertEqual(len(_tables(doc)), 1, "Expected exactly one DOCX table")

    def test_table_dimensions(self):
        doc = _convert(self.SIMPLE_TABLE)
        table = _tables(doc)[0]
        self.assertEqual(len(table.rows), 3)  # header + 2 data rows
        self.assertEqual(len(table.columns), 2)

    def test_header_row_is_bold(self):
        doc = _convert(self.SIMPLE_TABLE)
        table = _tables(doc)[0]
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    self.assertTrue(run.bold, f"Header cell '{run.text}' should be bold")

    def test_data_rows_not_bold(self):
        doc = _convert(self.SIMPLE_TABLE)
        table = _tables(doc)[0]
        for row in table.rows[1:]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        self.assertFalse(run.bold, f"Data cell '{run.text}' should not be bold")

    def test_cell_text_content(self):
        doc = _convert(self.SIMPLE_TABLE)
        table = _tables(doc)[0]
        self.assertEqual(table.cell(0, 0).text, "Name")
        self.assertEqual(table.cell(0, 1).text, "Age")
        self.assertEqual(table.cell(1, 0).text, "Alice")
        self.assertEqual(table.cell(1, 1).text, "30")
        self.assertEqual(table.cell(2, 0).text, "Bob")
        self.assertEqual(table.cell(2, 1).text, "25")

    def test_no_extra_paragraphs_for_table_lines(self):
        """Table lines must NOT produce standalone paragraphs."""
        doc = _convert(self.SIMPLE_TABLE)
        # All paragraphs come from inside table cells, not from the body
        body_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        # python-docx includes an empty paragraph after a table; ignore empty ones
        self.assertEqual(len(body_paragraphs), 0,
                         f"Unexpected body paragraphs: {[p.text for p in body_paragraphs]}")


class TestTableWithoutOuterPipes(unittest.TestCase):
    TABLE_NO_OUTER = "Name | Age\n---|---\nAlice | 30"

    def test_table_created(self):
        doc = _convert(self.TABLE_NO_OUTER)
        self.assertEqual(len(_tables(doc)), 1)

    def test_cell_text(self):
        doc = _convert(self.TABLE_NO_OUTER)
        table = _tables(doc)[0]
        self.assertEqual(table.cell(0, 0).text, "Name")
        self.assertEqual(table.cell(1, 0).text, "Alice")


class TestTableAlignment(unittest.TestCase):
    TABLE = "| L | R | C |\n|:---|---:|:---:|\n| a | b | c |"

    def test_alignments(self):
        doc = _convert(self.TABLE)
        table = _tables(doc)[0]
        # Header row
        self.assertEqual(table.cell(0, 0).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)
        self.assertEqual(table.cell(0, 1).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.RIGHT)
        self.assertEqual(table.cell(0, 2).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
        # Data row
        self.assertEqual(table.cell(1, 0).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)
        self.assertEqual(table.cell(1, 1).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.RIGHT)
        self.assertEqual(table.cell(1, 2).paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)


class TestTableEscapedPipes(unittest.TestCase):
    def test_escaped_pipe_in_cell(self):
        md = r"| A\|B | C |" + "\n|---|---|\n" + r"| x\|y | z |"
        doc = _convert(md)
        table = _tables(doc)[0]
        self.assertEqual(table.cell(0, 0).text, "A|B")
        self.assertEqual(table.cell(1, 0).text, "x|y")


class TestTablePaddingShortRows(unittest.TestCase):
    def test_short_row_padded(self):
        md = "| A | B | C |\n|---|---|---|\n| only_one |"
        doc = _convert(md)
        table = _tables(doc)[0]
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(table.cell(1, 0).text, "only_one")
        self.assertEqual(table.cell(1, 1).text, "")
        self.assertEqual(table.cell(1, 2).text, "")


class TestTableDoesNotBreakOtherElements(unittest.TestCase):
    def test_heading_before_table(self):
        md = "# Title\n\n| A | B |\n|---|---|\n| 1 | 2 |"
        doc = _convert(md)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertEqual(len(headings), 1)
        self.assertEqual(headings[0].text, "Title")
        self.assertEqual(len(_tables(doc)), 1)

    def test_paragraph_after_table(self):
        md = "| A | B |\n|---|---|\n| 1 | 2 |\n\nFinal paragraph."
        doc = _convert(md)
        self.assertEqual(len(_tables(doc)), 1)
        body_paras = [p for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any("Final paragraph" in p.text for p in body_paras))

    def test_code_block_not_treated_as_table(self):
        md = "```\n| fake | table |\n|---|---|\n```"
        doc = _convert(md)
        self.assertEqual(len(_tables(doc)), 0, "Code block content must not become a table")

    def test_table_at_eof_no_trailing_newline(self):
        md = "| X | Y |\n|---|---|\n| 1 | 2 |"
        doc = _convert(md)
        self.assertEqual(len(_tables(doc)), 1)


class TestTableBorders(unittest.TestCase):
    TABLE = "| A | B |\n|---|---|\n| 1 | 2 |"

    def _get_tbl_borders(self, doc):
        table = doc.tables[0]
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            return None
        return tblPr.find(qn('w:tblBorders'))

    def test_default_border_color_is_gray(self):
        """By default the table must have single-line gray (808080) borders."""
        converter = MarkdownToDocx()
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        self.assertIsNotNone(borders, "w:tblBorders element must be present")
        top = borders.find(qn('w:top'))
        self.assertIsNotNone(top)
        self.assertEqual(top.get(qn('w:val')), 'single')
        self.assertEqual(top.get(qn('w:color')).upper(), '808080')

    def test_no_borders_flag_removes_visible_borders(self):
        """With table_borders=False every border side must be 'none'."""
        converter = MarkdownToDocx(table_borders=False)
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        self.assertIsNotNone(borders, "w:tblBorders element must still be present")
        for side in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
            el = borders.find(qn(side))
            self.assertIsNotNone(el, f"{side} must be present")
            self.assertEqual(el.get(qn('w:val')), 'none', f"{side} must be 'none'")

    def test_all_six_border_sides_present_when_visible(self):
        """All six border directions must be set when borders are visible."""
        converter = MarkdownToDocx(table_borders=True)
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        for side in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
            self.assertIsNotNone(borders.find(qn(side)), f"{side} must be present")

    def test_default_constructor_has_borders(self):
        """MarkdownToDocx() with no arguments defaults to single-line borders."""
        converter = MarkdownToDocx()
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        self.assertIsNotNone(borders)
        top = borders.find(qn('w:top'))
        self.assertEqual(top.get(qn('w:val')), 'single')

    def test_custom_border_color_hex(self):
        """A custom hex color is applied to all six border sides."""
        converter = MarkdownToDocx(table_border_color='FF0000')
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        for side in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
            el = borders.find(qn(side))
            self.assertEqual(el.get(qn('w:color')).upper(), 'FF0000',
                             f"{side} color must be FF0000")

    def test_custom_border_color_rgb(self):
        """A custom R,G,B color string is correctly converted and applied."""
        converter = MarkdownToDocx(table_border_color='0,0,255')
        doc = converter.convert(self.TABLE)
        borders = self._get_tbl_borders(doc)
        top = borders.find(qn('w:top'))
        self.assertEqual(top.get(qn('w:color')).upper(), '0000FF')


class TestParseColor(unittest.TestCase):
    def test_hex_lowercase(self):
        self.assertEqual(_parse_color('808080'), '808080')

    def test_hex_uppercase(self):
        self.assertEqual(_parse_color('FF0000'), 'FF0000')

    def test_hex_with_hash(self):
        self.assertEqual(_parse_color('#0000ff'), '0000FF')

    def test_rgb_integers(self):
        self.assertEqual(_parse_color('128,128,128'), '808080')

    def test_rgb_black(self):
        self.assertEqual(_parse_color('0,0,0'), '000000')

    def test_rgb_white(self):
        self.assertEqual(_parse_color('255,255,255'), 'FFFFFF')

    def test_invalid_hex_raises(self):
        with self.assertRaises(ValueError):
            _parse_color('ZZZZZZ')

    def test_invalid_rgb_range_raises(self):
        with self.assertRaises(ValueError):
            _parse_color('256,0,0')

    def test_invalid_rgb_count_raises(self):
        with self.assertRaises(ValueError):
            _parse_color('1,2')


class TestNoPreferredWidths(unittest.TestCase):
    TABLE = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |"

    def _get_table(self, md: str):
        converter = MarkdownToDocx()
        doc = converter.convert(md)
        return doc.tables[0]

    def test_no_table_preferred_width(self):
        """w:tblW must not be present in w:tblPr."""
        table = self._get_table(self.TABLE)
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            self.assertIsNone(tblPr.find(qn('w:tblW')),
                              "w:tblW (table preferred width) must not be set")

    def test_no_column_preferred_width(self):
        """w:gridCol elements must not carry a w:w attribute."""
        table = self._get_table(self.TABLE)
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gridCol in tblGrid.findall(qn('w:gridCol')):
                self.assertNotIn(qn('w:w'), gridCol.attrib,
                                 "w:gridCol must not have a w:w attribute (column preferred width)")

    def test_no_cell_preferred_width(self):
        """w:tcW must not be present inside any cell's w:tcPr."""
        table = self._get_table(self.TABLE)
        for tc in table._element.iter(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                self.assertIsNone(tcPr.find(qn('w:tcW')),
                                  "w:tcW (cell preferred width) must not be set")


if __name__ == "__main__":
    unittest.main()
