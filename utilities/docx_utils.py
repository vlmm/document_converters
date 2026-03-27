import argparse
import copy
import os

from docx import Document
from docx.enum.text import WD_BREAK


from docx.enum.text import WD_BREAK

def merge_docx(files, output):
    merged_document = Document()
    # Премахваме първия implicit параграф, за да няма празни параграфи в началото
    merged_document._body.clear_content()

    first_file = True
    for file in files:
        sub_doc = Document(file)

        if not first_file:
            # Добавяме page break към последния параграф на обединения документ
            # Ако няма параграфи, създаваме нов параграф
            if merged_document.paragraphs:
                last_paragraph = merged_document.paragraphs[-1]
            else:
                last_paragraph = merged_document.add_paragraph()
            run = last_paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
        else:
            first_file = False

        # Копираме всички елементи на body от sub_doc в merged_document
        for element in sub_doc.element.body:
            merged_document.element.body.append(copy.deepcopy(element))

    merged_document.save(output)
    print(f"Merged {len(files)} files into {output}")


def split_docx_by_sections(input_file, output_prefix):
    doc = Document(input_file)
    num_sections = len(doc.sections)
    print(f"Document has {num_sections} sections.")
    
    # Създава по един файл за всяка секция
    for i in range(num_sections):
        new_doc = Document()
        new_doc._body.clear_content()

        # Към всяка секция добавяме параграфи и таблици, които са в тази секция
        # В python-docx няма директен API за параграфите по секции.
        # Решение: работим с елементите на body и разделяме по секции чрез манипулиране на XML
        # Алтернативно, може да разделим документа на части според sectionProperties в XML, 
        # но тук ще ползваме подход с копиране на конкретни части от XML.

        # Този метод е опростен и може да пропуска някои елементи,
        # но при тестове дава базово разделяне.

        # ДОБАВЯМЕ САМО ТЕКУЩАТА СЕКЦИЯ (чрез XML отделяне)
        body = doc.element.body
        new_body = new_doc.element.body

        # Събиране на елементи в настоящия документ
        # Трябва да отделим елементите от XML за секциите. 
        # findall и след това съобразно секция (важно)

        children = list(body.iterchildren())
        # търсим секция (sectPr) елементи, с които започва секция
        
        sect_positions = []
        for idx, el in enumerate(children):
            sectPr = el.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr")
            if sectPr is not None:
                sect_positions.append(idx)

        # Добавяме елементите от начало до първата sectPr за първата секция, между sectPr за другите и т.н.
        start_idx = sect_positions[i-1]+1 if i > 0 else 0
        end_idx = sect_positions[i] if i < len(sect_positions) else len(children)

        # Копираме елементите от start_idx до end_idx в новия документ
        for el in children[start_idx:end_idx]:
            new_body.append(copy.deepcopy(el))

        # Винаги трябва да добавим sectPr за тази секция, за да се запази форматирането
        # След end_idx добавяме sectPr от позицията sect_positions[i] ако има такава
        if i < len(sect_positions):
            new_body.append(copy.deepcopy(children[sect_positions[i]]))

        out_fname = f"{output_prefix}_{i+1}.docx"
        new_doc.save(out_fname)
        print(f"Saved section {i+1} to {out_fname}")


def main():
    parser = argparse.ArgumentParser(description="Merge or split DOCX files")
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")

    # Merge
    parser_merge = subparsers.add_parser("merge", help="Merge DOCX files")
    parser_merge.add_argument("input_files", nargs="+", help="Input DOCX files to merge")
    parser_merge.add_argument("-o", "--output", required=True, help="Output merged DOCX filename")

    # Split
    parser_split = subparsers.add_parser("split", help="Split DOCX by sections")
    parser_split.add_argument("input_file", help="Input DOCX file to split")
    parser_split.add_argument("-o", "--output_prefix", required=True, help="Output files prefix")

    args = parser.parse_args()

    if args.command == "merge":
        merge_docx(args.input_files, args.output)
    elif args.command == "split":
        split_docx_by_sections(args.input_file, args.output_prefix)
    else:
        print("No valid command. Use 'merge' or 'split'.")


if __name__ == "__main__":
    main()